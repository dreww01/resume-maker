"""Unit and integration tests for FastAPI backend routes and Pydantic schemas."""

import io
from unittest.mock import patch
import pytest
from docx import Document
from fastapi.testclient import TestClient

from src.api import app
from src.database import create_resume, update_resume

client = TestClient(app)


def test_root_endpoint():
    """Verify GET / returns 200 and expected welcome message schema."""
    response = client.get("/")
    assert response.status_code == 200
    data = response.json()
    assert "message" in data
    assert "Resume Tailor API" in data["message"]


def test_upload_invalid_file_extension():
    """Verify uploading non-PDF/DOCX file returns 400 Bad Request."""
    response = client.post(
        "/upload",
        files={"file": ("resume.txt", b"plain text", "text/plain")}
    )
    assert response.status_code == 400
    assert response.json()["detail"] == "File must be .pdf or .docx"


def test_upload_empty_file():
    """Verify uploading empty file returns 400 Bad Request."""
    response = client.post(
        "/upload",
        files={"file": ("empty.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 400
    assert "empty" in response.json()["detail"].lower()


def test_upload_valid_docx():
    """Verify uploading valid DOCX returns 201 Created with typed schema attributes."""
    doc = Document()
    doc.add_paragraph("Alice Smith")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    response = client.post(
        "/upload",
        files={"file": ("alice_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 201
    data = response.json()
    assert "id" in data
    assert data["filename"] == "alice_resume.docx"
    assert data["status"] == "uploaded"
    assert "created_at" in data


def test_get_resume_status():
    """Verify GET /resumes/{id} returns typed ResumeResponse schema."""
    resume_id = create_resume("status_test.docx", b"content")
    response = client.get(f"/resumes/{resume_id}")
    assert response.status_code == 200
    data = response.json()
    assert data["id"] == resume_id
    assert data["filename"] == "status_test.docx"
    assert data["original_filename"] == "status_test.docx"
    assert data["status"] == "uploaded"
    assert data["has_tailored_resume"] is False
    assert data["has_output"] is False
    assert data["has_cover_letter"] is False
    assert data["created_at"] is not None


def test_get_nonexistent_resume_status():
    """Verify requesting nonexistent resume ID returns 404 Not Found."""
    response = client.get("/resumes/999999")
    assert response.status_code == 404
    assert response.json()["detail"] == "Resume not found"


@patch("src.api.call_openai")
def test_tailor_resume_flow_with_plain_text(mock_call_openai):
    """Verify tailoring flow using plain text request body."""
    doc = Document()
    doc.add_paragraph("Bob Jones")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    resume_id = create_resume("bob_resume.docx", docx_bytes)

    mock_call_openai.return_value = {
        "name": "Bob Jones",
        "email": "bob@example.com",
        "phone": "+1 555-0199",
        "location": "New York, NY",
        "professional_summary": "Tailored engineer summary",
        "work_experience": [],
        "projects": [],
        "skills": ["Python", "FastAPI"],
        "soft_skills": ["Teamwork"],
        "education": []
    }

    response = client.post(
        f"/resumes/{resume_id}/tailor",
        content="We need a Python developer with FastAPI skills",
        headers={"Content-Type": "text/plain"}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "completed"
    assert data["user_name"] == "Bob Jones"

    # Verify download of tailored resume
    dl_response = client.get(f"/resumes/{resume_id}/download")
    assert dl_response.status_code == 200
    assert "Bob_Jones_resume_" in dl_response.headers.get("content-disposition", "")
    assert dl_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@patch("src.api.call_openai")
def test_tailor_resume_flow_with_json_schema(mock_call_openai):
    """Verify tailoring flow using structured JSON request body."""
    doc = Document()
    doc.add_paragraph("Bob Jones")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    resume_id = create_resume("bob_resume_json.docx", docx_bytes)

    mock_call_openai.return_value = {
        "name": "Bob Jones",
        "email": "bob@example.com",
        "phone": "+1 555-0199",
        "location": "New York, NY",
        "professional_summary": "Tailored engineer summary",
        "work_experience": [],
        "projects": [],
        "skills": ["Python", "FastAPI"],
        "soft_skills": ["Teamwork"],
        "education": []
    }

    response = client.post(
        f"/resumes/{resume_id}/tailor",
        json={"job_description": "We need a Senior Python developer with FastAPI and Kubernetes skills."}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "completed"
    assert data["user_name"] == "Bob Jones"


def test_tailor_resume_short_description_fails():
    """Verify job description with <10 characters returns 400 Bad Request."""
    resume_id = create_resume("short_desc.docx", b"dummy content")
    response = client.post(
        f"/resumes/{resume_id}/tailor",
        json={"job_description": "short"}
    )
    assert response.status_code == 400


@patch("src.api.call_openai")
def test_tailor_resume_ai_upstream_failure(mock_call_openai):
    """Verify 502 Bad Gateway is returned when upstream AI call raises ValueError."""
    doc = Document()
    doc.add_paragraph("Bob Jones")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    resume_id = create_resume("ai_fail.docx", docx_bytes)
    mock_call_openai.side_effect = ValueError("OpenAI rate limit exceeded or invalid JSON")

    response = client.post(
        f"/resumes/{resume_id}/tailor",
        json={"job_description": "We need a Senior Python developer with FastAPI and Kubernetes skills."}
    )
    assert response.status_code == 502
    assert "upstream" in response.json()["detail"].lower() or "ai service failure" in response.json()["detail"].lower()


@patch("src.api.call_openai_cover_letter")
def test_generate_cover_letter_flow(mock_call_openai_cover_letter):
    """Verify cover letter generation flow and download endpoint."""
    doc = Document()
    doc.add_paragraph("Bob Jones")
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    resume_id = create_resume("bob_cover.docx", docx_bytes)

    mock_call_openai_cover_letter.return_value = {
        "name": "Bob Jones",
        "content": "Dear Hiring Manager,\n\nI am thrilled to apply."
    }

    response = client.post(
        f"/resumes/{resume_id}/cover-letter",
        content="Job description for Python backend position",
        headers={"Content-Type": "text/plain"}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "completed"
    assert data["user_name"] == "Bob Jones"

    # Verify download of cover letter
    dl_response = client.get(f"/resumes/{resume_id}/cover-letter/download")
    assert dl_response.status_code == 200
    assert "Bob_Jones_cover_letter_" in dl_response.headers.get("content-disposition", "")
    assert dl_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_download_resume_not_ready():
    """Verify downloading pending resume returns 400."""
    resume_id = create_resume("pending.docx", b"bytes")
    response = client.get(f"/resumes/{resume_id}/download")
    assert response.status_code == 400
    assert response.json()["detail"] == "Resume not ready for download"


def test_download_nonexistent_resume():
    """Verify downloading nonexistent resume returns 404."""
    response = client.get("/resumes/999999/download")
    assert response.status_code == 404
