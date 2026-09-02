"""AI resume processor and DOCX builder engine.

Handles multimodal vision PDF text extraction, OpenAI chat completions for tailoring
and cover letter synthesis, defensive output validation via Pydantic schemas, and structured
DOCX document styling and generation.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import tempfile
from typing import Any, Optional

import docx
from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI
from openai.types.chat import ChatCompletionMessageParam
from pydantic import ValidationError

from src.prompts.cover_letter import COVER_LETTER_SYSTEM_PROMPT, COVER_LETTER_USER_TEMPLATE
from src.prompts.resume_tailor import SYSTEM_PROMPT, USER_PROMPT_TEMPLATE
from src.schemas import CoverLetterData, TailoredResumeData

load_dotenv(override=True)

logger = logging.getLogger(__name__)

OPENAI_API_KEY: Optional[str] = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL: Optional[str] = os.getenv("OPENAI_BASE_URL")
AI_MODEL: str = os.getenv("AI_MODEL", "gpt-4o-mini")
VISION_MODEL: str = os.getenv("VISION_MODEL", "gpt-4o-mini")

# Styling Constants
PRIMARY_COLOR = RGBColor(0, 51, 102)      # Deep Navy for headings
MUTED_TEXT_COLOR = RGBColor(64, 64, 64)   # Gray for secondary info
DARK_TEXT_COLOR = RGBColor(0, 0, 0)       # Black for main text


def get_openai_client() -> OpenAI:
    """Instantiate and return a configured OpenAI client.

    Returns:
        Configured OpenAI client instance.

    Raises:
        ValueError: If OPENAI_API_KEY is not configured in the environment.
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not set. Add it to your .env file.")

    base_url = os.getenv("OPENAI_BASE_URL")
    kwargs: dict[str, Any] = {"api_key": api_key}
    if base_url:
        kwargs["base_url"] = base_url
    return OpenAI(**kwargs)


# ---------------------------------------------------------------------------
# DOCX Document Builder Helper Functions
# ---------------------------------------------------------------------------

def set_document_margins(
    doc: _DocumentType,
    top: float = 0.5,
    bottom: float = 0.5,
    left: float = 0.5,
    right: float = 0.5,
) -> None:
    """Set uniform margins across all sections of a Word document.

    Args:
        doc: The python-docx Document object to configure.
        top: Top margin in inches.
        bottom: Bottom margin in inches.
        left: Left margin in inches.
        right: Right margin in inches.
    """
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_styled_heading(
    doc: _DocumentType,
    text: str,
    font_size: float = 11.0,
    color: RGBColor = PRIMARY_COLOR,
    space_before: float = 10.0,
    space_after: float = 2.0,
) -> None:
    """Append a standardized uppercase section heading to a document.

    Args:
        doc: The python-docx Document object.
        text: Heading text content.
        font_size: Font size in points.
        color: Font RGB color.
        space_before: Paragraph spacing before heading in points.
        space_after: Paragraph spacing after heading in points.
    """
    heading = doc.add_paragraph()
    heading_run = heading.add_run(text.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(font_size)
    heading_run.font.color.rgb = color
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)


def add_bullet_item(
    doc: _DocumentType,
    text: str,
    space_after: float = 2.0,
) -> None:
    """Append a styled bullet list item to a document.

    Args:
        doc: The python-docx Document object.
        text: Bullet text content.
        space_after: Spacing after paragraph in points.
    """
    bullet_para = doc.add_paragraph(text, style="List Bullet")
    bullet_para.paragraph_format.space_after = Pt(space_after)


# ---------------------------------------------------------------------------
# PDF & Document Extraction
# ---------------------------------------------------------------------------

def extract_pdf_with_vision(file_bytes: bytes) -> str:
    """Convert PDF pages to images and extract text using multimodal Vision AI.

    Args:
        file_bytes: Raw binary bytes of the PDF file.

    Returns:
        Clean textual content extracted from all PDF pages.

    Raises:
        ValueError: If PDF rendering or API extraction fails.
    """
    import pypdfium2 as pdfium  # type: ignore[import-untyped]

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        pdf = pdfium.PdfDocument(tmp_path)
        image_contents: list[dict[str, Any]] = []

        for page in pdf:
            bitmap = page.render(scale=2)
            pil_image = bitmap.to_pil()
            buffer = io.BytesIO()
            pil_image.save(buffer, format="PNG")
            base64_image = base64.b64encode(buffer.getvalue()).decode("utf-8")
            image_contents.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{base64_image}"},
            })
        pdf.close()
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    client = get_openai_client()
    vision_model = os.getenv("VISION_MODEL", VISION_MODEL)
    messages: list[Any] = [{
        "role": "user",
        "content": [
            {
                "type": "text",
                "text": "Extract ALL text from this resume image. Preserve the structure and sections. Return only the extracted text, no commentary.",
            },
            *image_contents,
        ],
    }]

    response = client.chat.completions.create(
        model=vision_model,
        messages=messages,
        max_tokens=4000,
    )

    extracted = response.choices[0].message.content or ""
    return extracted.strip()


def read_resume(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from an uploaded resume file (PDF or DOCX).

    Args:
        file_bytes: Raw binary content of the file.
        filename: Name of the uploaded file to determine extraction strategy.

    Returns:
        Extracted plain text.

    Raises:
        ValueError: If the file format is unsupported or parsing fails.
    """
    if filename.lower().endswith(".pdf"):
        return extract_pdf_with_vision(file_bytes)

    if filename.lower().endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(paragraphs).strip()

    raise ValueError("File must be .pdf or .docx")


# ---------------------------------------------------------------------------
# AI Completion & Schema Validation
# ---------------------------------------------------------------------------

def call_openai(resume_text: str, job_description: str) -> dict[str, Any]:
    """Execute AI completion to tailor a resume and validate response structure.

    Args:
        resume_text: Plain text of candidate's original resume.
        job_description: Plain text of the target job description.

    Returns:
        Dictionary validated against TailoredResumeData schema.

    Raises:
        ValueError: If OpenAI returns invalid JSON or structure does not match schema.
    """
    client = get_openai_client()
    ai_model = os.getenv("AI_MODEL", AI_MODEL)
    user_prompt = USER_PROMPT_TEMPLATE.format(
        resume_text=resume_text,
        job_description=job_description,
    )

    messages: list[ChatCompletionMessageParam] = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    response = client.chat.completions.create(
        model=ai_model,
        messages=messages,
        response_format={"type": "json_object"},
    )

    raw_content = response.choices[0].message.content or "{}"
    try:
        parsed_json = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        logger.error("AI response was not valid JSON: %s", raw_content)
        raise ValueError(f"AI response was not valid JSON: {exc}") from exc

    try:
        validated_model = TailoredResumeData.model_validate(parsed_json)
        return validated_model.model_dump()
    except ValidationError as exc:
        logger.error("AI tailored resume output validation failed: %s", exc)
        raise ValueError(f"AI tailored resume output validation failed: {exc}") from exc


def call_openai_cover_letter(resume_text: str, job_description: str) -> dict[str, Any]:
    """Execute AI completion to generate a cover letter and validate response structure.

    Args:
        resume_text: Plain text of candidate's original resume.
        job_description: Plain text of the target job description.

    Returns:
        Dictionary validated against CoverLetterData schema containing 'content' and 'name'.

    Raises:
        ValueError: If OpenAI returns invalid JSON or structure is invalid.
    """
    client = get_openai_client()
    ai_model = os.getenv("AI_MODEL", AI_MODEL)
    user_prompt = COVER_LETTER_USER_TEMPLATE.format(
        resume_text=resume_text,
        job_description=job_description,
    )

    messages: list[ChatCompletionMessageParam] = [
        {"role": "system", "content": COVER_LETTER_SYSTEM_PROMPT},
        {"role": "user", "content": user_prompt},
    ]

    response = client.chat.completions.create(
        model=ai_model,
        messages=messages,
        response_format={"type": "json_object"},
    )

    raw_content = response.choices[0].message.content or "{}"
    try:
        parsed_json = json.loads(raw_content)
    except json.JSONDecodeError as exc:
        logger.error("AI response was not valid JSON: %s", raw_content)
        raise ValueError(f"AI response was not valid JSON: {exc}") from exc

    try:
        validated_model = CoverLetterData.model_validate(parsed_json)
        return validated_model.model_dump()
    except ValidationError as exc:
        logger.error("AI cover letter output validation failed: %s", exc)
        raise ValueError(f"AI cover letter output validation failed: {exc}") from exc


# ---------------------------------------------------------------------------
# Document Generation
# ---------------------------------------------------------------------------

def create_cover_letter_docx(text: str) -> bytes:
    """Generate a cleanly formatted .docx Word document containing the cover letter.

    Args:
        text: Body text of the cover letter with double newline paragraph separation.

    Returns:
        Raw bytes of generated .docx file.
    """
    doc = Document()
    set_document_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0)

    for paragraph in text.strip().split("\n\n"):
        clean_p = paragraph.strip()
        if clean_p:
            p = doc.add_paragraph(clean_p)
            p.paragraph_format.space_after = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_docx(resume_data: dict[str, Any] | TailoredResumeData) -> bytes:
    """Generate an ATS-compliant, beautifully styled .docx Word resume document.

    Args:
        resume_data: Structured resume data dictionary or TailoredResumeData instance.

    Returns:
        Raw bytes of generated .docx file.

    Raises:
        ValueError: If required data elements (e.g. name) are missing.
    """
    if isinstance(resume_data, TailoredResumeData):
        data = resume_data.model_dump()
    elif isinstance(resume_data, dict):
        data = resume_data
    else:
        raise ValueError("resume_data must be a dict or TailoredResumeData instance")

    doc = Document()
    set_document_margins(doc, top=0.5, bottom=0.5, left=0.5, right=0.5)

    # 1. Header: Candidate Name
    candidate_name = str(data.get("name") or "").strip() or "Resume"
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = DARK_TEXT_COLOR
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)

    # 2. Header: Contact Information
    contact_parts: list[str] = []
    for field in ["email", "phone", "location", "github", "linkedin", "portfolio"]:
        value = str(data.get(field) or "").strip()
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = MUTED_TEXT_COLOR
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(10)

    # 3. Professional Summary
    summary_text = str(data.get("professional_summary") or "").strip()
    if summary_text:
        add_styled_heading(doc, "Professional Summary")
        summary_para = doc.add_paragraph(summary_text)
        summary_para.paragraph_format.space_after = Pt(8)

    # 4. Professional Experience
    work_experience = data.get("work_experience") or []
    if work_experience:
        add_styled_heading(doc, "Professional Experience")
        for job in work_experience:
            if not isinstance(job, dict):
                continue
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(str(job.get("title") or ""))
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_para.paragraph_format.space_after = Pt(2)

            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{str(job.get('company') or '')} | ")
            company_run.font.size = Pt(10)
            duration_run = company_para.add_run(str(job.get("duration") or ""))
            duration_run.italic = True
            duration_run.font.size = Pt(10)
            duration_run.font.color.rgb = MUTED_TEXT_COLOR
            company_para.paragraph_format.space_after = Pt(4)

            bullets = job.get("bullets") or []
            for bullet in bullets:
                if bullet and str(bullet).strip():
                    add_bullet_item(doc, str(bullet).strip(), space_after=2.0)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    # 5. Key Projects
    projects = data.get("projects") or []
    if projects:
        add_styled_heading(doc, "Key Projects")
        for project in projects:
            if not isinstance(project, dict):
                continue
            project_para = doc.add_paragraph()
            proj_name_run = project_para.add_run(str(project.get("name") or ""))
            proj_name_run.bold = True
            proj_name_run.font.size = Pt(10.5)
            project_para.paragraph_format.space_after = Pt(2)

            bullets = project.get("bullets") or []
            for bullet in bullets:
                if bullet and str(bullet).strip():
                    add_bullet_item(doc, str(bullet).strip(), space_after=2.0)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)

    # 6. Technical Skills
    skills = data.get("skills") or []
    if skills:
        valid_skills = [str(s) for s in skills if s is not None and str(s).strip()]
        if valid_skills:
            add_styled_heading(doc, "Technical Skills")
            skills_para = doc.add_paragraph(" | ".join(valid_skills))
            skills_para.paragraph_format.space_after = Pt(8)

    # 7. Core Competencies / Soft Skills
    soft_skills = data.get("soft_skills") or []
    if soft_skills:
        valid_soft_skills = [str(s) for s in soft_skills if s is not None and str(s).strip()]
        if valid_soft_skills:
            add_styled_heading(doc, "Core Competencies")
            soft_skills_para = doc.add_paragraph(" | ".join(valid_soft_skills))
            soft_skills_para.paragraph_format.space_after = Pt(8)

    # 8. Education
    education = data.get("education") or []
    if education:
        add_styled_heading(doc, "Education")
        for edu in education:
            if not isinstance(edu, dict):
                continue
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(str(edu.get("degree") or ""))
            degree_run.bold = True
            degree_run.font.size = Pt(10.5)
            edu_para.add_run(f" - {str(edu.get('institution') or '')}, {str(edu.get('year') or '')}")
            edu_para.paragraph_format.space_after = Pt(4)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
