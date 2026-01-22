import os
import json
import PyPDF2
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

#Read resume from PDF or DOCX file.
def read_resume(file_path):
    
    if file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text.strip()

    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()

    else:
        raise ValueError("File must be .pdf or .docx")


def call_openai(resume_text, job_description):
    
    client = OpenAI(api_key=openai_api_key)

    system_prompt = """You are a professional resume writer specializing in ATS optimization.

Your task:
1. Extract the candidate's information from the provided resume
2. Tailor content to align with the target job description
3. Integrate relevant keywords from the job description into experience bullets, skills, and summary
4. Write in clear, direct language without filler words, buzzwords, or em-dashes
5. Use strong action verbs and quantify achievements where possible
6. Ensure the professional summary directly addresses what the role requires

Guidelines:
- Mirror terminology from the job description (e.g., if they say "CI/CD pipelines," use that exact phrase)
- Keep bullets concise and results-focused
- Prioritize skills and experiences that match the job requirements
- Maintain a professional, confident tone throughout

Return the data in this exact JSON format:
{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "github": "github link",
    "professional_summary": "2-3 sentences tailored to the job",
    "work_experience": [
        {
            "title": "Job Title",
            "company": "Company Name",
            "duration": "Start - End",
            "bullets": ["achievement 1", "achievement 2"]
        }
    ],
    "projects": [
        {
            "name": "Project Name",
            "bullets": ["achievement/feature 1", "achievement/feature 2", "achievement/feature 3"]
        }
    ],
    "skills": [
        "Skill 1",
        "Skill 2",
        "Skill 3"
    ],
    "soft_skills": [
        "Skill 1",
        "Skill 2",
        "Skill 3"
    ],
    "education": [
        {
            "degree": "Degree",
            "institution": "Institution",
            "year": "Year"
        }
    ]
}"""

    user_prompt = f"""Here is the resume:
{resume_text}

Here is the job description:
{job_description}

Extract all information and tailor the resume to match this job. Use keywords from the job description."""

    response = client.chat.completions.create(
        model="gpt-4.1-nano",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        response_format={"type": "json_object"}
    )

    resume_data = json.loads(response.choices[0].message.content)
    return resume_data


def create_docx(resume_data, output_path="tailored_resume.docx"):
    """Creates a modern ATS-friendly resume in industry-standard order."""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_section_heading(text):
        heading = doc.add_paragraph()
        heading_run = heading.add_run(text.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        heading.space_before = Pt(10)
        heading.space_after = Pt(2)

    # === HEADER ===
    name = doc.add_paragraph()
    name_run = name.add_run(resume_data['name'])
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(4)

    contact = doc.add_paragraph()
    contact_run = contact.add_run(
        f"{resume_data['email']} | {resume_data['phone']} | {resume_data['github']}"
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(64, 64, 64)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(10)

    # === 1. PROFESSIONAL SUMMARY ===
    add_section_heading('Professional Summary')
    summary = doc.add_paragraph(resume_data['professional_summary'])
    summary.space_after = Pt(8)

    # === 2. PROFESSIONAL EXPERIENCE ===
    add_section_heading('Professional Experience')
    for job in resume_data['work_experience']:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(job['title'])
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_para.space_after = Pt(2)

        company_para = doc.add_paragraph()
        company_run = company_para.add_run(f"{job['company']} | ")
        company_run.font.size = Pt(10)
        duration_run = company_para.add_run(job['duration'])
        duration_run.italic = True
        duration_run.font.size = Pt(10)
        duration_run.font.color.rgb = RGBColor(64, 64, 64)
        company_para.space_after = Pt(4)

        for bullet in job['bullets']:
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.space_after = Pt(2)

        doc.add_paragraph().space_after = Pt(6)

    # === 3. KEY PROJECTS ===
    add_section_heading('Key Projects')
    for project in resume_data['projects']:
        project_para = doc.add_paragraph()
        name_run = project_para.add_run(project['name'])
        name_run.bold = True
        name_run.font.size = Pt(10.5)
        project_para.space_after = Pt(2)

        for bullet in project['bullets']:
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.space_after = Pt(2)

        doc.add_paragraph().space_after = Pt(4)

    # === 4. TECHNICAL SKILLS ===
    add_section_heading('Technical Skills')
    skills_text = ' • '.join(resume_data['skills'])
    skills = doc.add_paragraph(skills_text)
    skills.space_after = Pt(8)

    # === 5. CORE COMPETENCIES ===
    add_section_heading('Core Competencies')
    soft_skills_text = ' • '.join(resume_data['soft_skills'])
    competencies = doc.add_paragraph(soft_skills_text)
    competencies.space_after = Pt(8)

    # === 6. EDUCATION ===
    add_section_heading('Education')
    for edu in resume_data.get('education', []):
        edu_para = doc.add_paragraph()
        degree_run = edu_para.add_run(edu['degree'])
        degree_run.bold = True
        degree_run.font.size = Pt(10.5)
        edu_para.add_run(f" — {edu['institution']}, {edu['year']}")
        edu_para.space_after = Pt(4)

    doc.save(output_path)
    return output_path
